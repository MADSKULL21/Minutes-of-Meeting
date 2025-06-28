import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import time
from datetime import datetime
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from SpeechToText import SpeechRecognition, SetAssistantStatus

class SpeechToTextGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Speech to Text Recorder")
        self.root.geometry("600x500")
        self.root.configure(bg='#f0f0f0')
        
        # Variables
        self.is_recording = False
        self.is_paused = False
        self.recorded_text = ""
        self.recording_thread = None
        
        self.setup_ui()
        
    def setup_ui(self):
        # Main frame
        main_frame = tk.Frame(self.root, bg='#f0f0f0')
        main_frame.pack(padx=20, pady=20, fill='both', expand=True)
        
        # Title
        title_label = tk.Label(main_frame, text="Speech to Text Recorder", 
                              font=('Arial', 18, 'bold'), bg='#f0f0f0', fg='#333')
        title_label.pack(pady=(0, 20))
        
        # Status frame
        status_frame = tk.Frame(main_frame, bg='#f0f0f0')
        status_frame.pack(fill='x', pady=(0, 20))
        
        self.status_label = tk.Label(status_frame, text="Status: Ready", 
                                    font=('Arial', 12), bg='#f0f0f0', fg='#666')
        self.status_label.pack()
        
        # Buttons frame
        buttons_frame = tk.Frame(main_frame, bg='#f0f0f0')
        buttons_frame.pack(pady=(0, 20))
        
        # Start button
        self.start_btn = tk.Button(buttons_frame, text="Start Recording", 
                                  command=self.start_recording,
                                  font=('Arial', 12), bg='#4CAF50', fg='white',
                                  width=15, height=2)
        self.start_btn.pack(side='left', padx=5)
        
        # Pause button
        self.pause_btn = tk.Button(buttons_frame, text="Pause Recording", 
                                  command=self.pause_recording,
                                  font=('Arial', 12), bg='#FF9800', fg='white',
                                  width=15, height=2, state='disabled')
        self.pause_btn.pack(side='left', padx=5)
        
        # Stop button
        self.stop_btn = tk.Button(buttons_frame, text="Stop Recording", 
                                 command=self.stop_recording,
                                 font=('Arial', 12), bg='#f44336', fg='white',
                                 width=15, height=2, state='disabled')
        self.stop_btn.pack(side='left', padx=5)
        
        # Export frame
        export_frame = tk.Frame(main_frame, bg='#f0f0f0')
        export_frame.pack(pady=(0, 20))
        
        # Export buttons
        self.export_word_btn = tk.Button(export_frame, text="Export to Word", 
                                        command=self.export_to_word,
                                        font=('Arial', 12), bg='#2196F3', fg='white',
                                        width=15, height=2, state='disabled')
        self.export_word_btn.pack(side='left', padx=5)
        
        self.export_pdf_btn = tk.Button(export_frame, text="Export to PDF", 
                                       command=self.export_to_pdf,
                                       font=('Arial', 12), bg='#9C27B0', fg='white',
                                       width=15, height=2, state='disabled')
        self.export_pdf_btn.pack(side='left', padx=5)
        
        # Text display frame
        text_frame = tk.Frame(main_frame, bg='#f0f0f0')
        text_frame.pack(fill='both', expand=True)
        
        # Text label
        text_label = tk.Label(text_frame, text="Recorded Text:", 
                             font=('Arial', 12, 'bold'), bg='#f0f0f0', fg='#333')
        text_label.pack(anchor='w', pady=(0, 5))
        
        # Text area with scrollbar
        text_container = tk.Frame(text_frame, bg='white', relief='sunken', bd=1)
        text_container.pack(fill='both', expand=True)
        
        self.text_area = tk.Text(text_container, wrap='word', font=('Arial', 11),
                                 bg='white', fg='#333', relief='flat')
        scrollbar = tk.Scrollbar(text_container, orient='vertical', command=self.text_area.yview)
        self.text_area.configure(yscrollcommand=scrollbar.set)
        
        self.text_area.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
    def start_recording(self):
        if not self.is_recording:
            self.is_recording = True
            self.is_paused = False
            self.status_label.config(text="Status: Recording...")
            self.start_btn.config(state='disabled')
            self.pause_btn.config(state='normal')
            self.stop_btn.config(state='normal')
            
            # Start recording in a separate thread
            self.recording_thread = threading.Thread(target=self.recording_loop)
            self.recording_thread.daemon = True
            self.recording_thread.start()
    
    def pause_recording(self):
        if self.is_recording:
            if self.is_paused:
                self.is_paused = False
                self.status_label.config(text="Status: Recording...")
                self.pause_btn.config(text="Resume Recording")
            else:
                self.is_paused = True
                self.status_label.config(text="Status: Paused")
                self.pause_btn.config(text="Pause Recording")
    
    def stop_recording(self):
        self.is_recording = False
        self.is_paused = False
        self.status_label.config(text="Status: Ready")
        self.start_btn.config(state='normal')
        self.pause_btn.config(state='disabled', text="Pause Recording")
        self.stop_btn.config(state='disabled')
        
        if self.recorded_text.strip():
            self.export_word_btn.config(state='normal')
            self.export_pdf_btn.config(state='normal')
    
    def recording_loop(self):
        while self.is_recording:
            if not self.is_paused:
                try:
                    # Get text from speech recognition
                    text = SpeechRecognition()
                    if text and text.strip():
                        self.recorded_text += text + " "
                        # Update GUI in main thread
                        self.root.after(0, self.update_text_area)
                except Exception as e:
                    print(f"Error in recording: {e}")
                    break
            time.sleep(0.1)
    
    def update_text_area(self):
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(1.0, self.recorded_text)
    
    def export_to_word(self):
        if not self.recorded_text.strip():
            messagebox.showwarning("Warning", "No text to export!")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            title="Save Word Document"
        )
        
        if file_path:
            try:
                doc = Document()
                doc.add_heading('Speech to Text Recording', 0)
                doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph('')
                doc.add_paragraph(self.recorded_text)
                doc.save(file_path)
                messagebox.showinfo("Success", f"Document saved to {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save document: {str(e)}")
    
    def export_to_pdf(self):
        if not self.recorded_text.strip():
            messagebox.showwarning("Warning", "No text to export!")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            title="Save PDF Document"
        )
        
        if file_path:
            try:
                doc = SimpleDocTemplate(file_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Title
                title = Paragraph('Speech to Text Recording', styles['Title'])
                story.append(title)
                story.append(Spacer(1, 12))
                
                # Date
                date_text = f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
                date_para = Paragraph(date_text, styles['Normal'])
                story.append(date_para)
                story.append(Spacer(1, 12))
                
                # Content
                content_para = Paragraph(self.recorded_text, styles['Normal'])
                story.append(content_para)
                
                doc.build(story)
                messagebox.showinfo("Success", f"PDF saved to {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save PDF: {str(e)}")

def main():
    root = tk.Tk()
    app = SpeechToTextGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()